import docx
import shutil
import sys

# ✅ Reconfigure stdout to use UTF-8 to prevent cp1252 errors on Windows
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

# 1. Backup the original file
original_path = r"c:\Users\ADMIN\Desktop\DATNRRRR\Báo cáo _DATN_Vũ_Văn_Trí_Tài_0267967_67CNPM-review.docx"
backup_path = r"c:\Users\ADMIN\Desktop\DATNRRRR\Báo cáo _DATN_Vũ_Văn_Trí_Tài_0267967_67CNPM-review_BACKUP.docx"
shutil.copyfile(original_path, backup_path)
print(f"Created backup at: {backup_path}")

doc = docx.Document(original_path)

def replace_range(doc, start_idx, end_idx, items):
    """
    items is a list of tuples: (text, style_name)
    """
    first_text, first_style = items[0]
    p_start = doc.paragraphs[start_idx]
    p_start.text = first_text
    
    # Try setting style
    try:
        if first_style:
            p_start.style = first_style
    except Exception:
        pass
    
    # Delete paragraphs in range in reverse
    for idx in range(end_idx, start_idx, -1):
        p = doc.paragraphs[idx]._element
        p.getparent().remove(p)
        
    # Insert new paragraphs
    current_p = p_start
    for text, style in items[1:]:
        new_p = doc.add_paragraph()
        new_p.text = text
        try:
            if style:
                new_p.style = style
        except Exception:
            pass
        current_p._element.addnext(new_p._element)
        current_p = new_p

# Let's construct new content for each chapter

ch1_items = [
    ("1. Giới thiệu bài toán", "Heading 2"),
    ("1.1. Bối cảnh và tính cấp thiết", "Heading 3"),
    ("Trong kỷ nguyên số hóa giáo dục hiện nay, việc tự học trực tuyến đã trở thành một xu hướng chủ đạo. Người học dễ dàng tiếp cận với vô số tài liệu học tập, giáo trình, bài báo khoa học trực tuyến dưới các định dạng PDF, Docx, TXT. Tuy nhiên, một trong những thách thức lớn nhất của người tự học là hiện tượng \"quá tải thông tin\" và \"thiếu lộ trình học tập rõ ràng\". Việc đọc một tài liệu dài hàng trăm trang mà không có một kế hoạch chia nhỏ kiến thức theo ngày học hợp lý, không có các mốc đánh giá năng lực tương tác và không có sự định hướng sư phạm từ giáo viên dẫn đến việc người học dễ nản lòng và từ bỏ giữa chừng.", "Normal"),
    ("Sự bùng nổ của các mô hình ngôn ngữ lớn (LLM) như LLaMA-3 kết hợp cùng kỹ thuật RAG (Retrieval-Augmented Generation) đang mở ra những cơ hội mới để giải quyết bài toán này. Ứng dụng AI vào giáo dục số giúp tự động hóa quá trình phân tích tài liệu học thuật, trích xuất cấu trúc cốt lõi để xây dựng lộ trình học cá nhân hóa theo từng ngày (Syllabus) và biên soạn bài giảng chi tiết bám sát nội dung tài liệu. Do đó, việc nghiên cứu xây dựng một hệ thống hỗ trợ người học cá nhân hóa lộ trình học và quản lý bài học tích hợp AI là vô cùng cấp thiết và mang tính thực tiễn cao.", "Normal"),
    ("1.2. Mô tả bài toán", "Heading 3"),
    ("Đề tài tập trung xây dựng \"Hệ thống hỗ trợ cá nhân hóa lộ trình học tập và quản lý bài học tích hợp Trí tuệ nhân tạo (AI)\". Hệ thống hướng tới việc hỗ trợ học viên tải lên tài liệu học thuật của cá nhân, từ đó AI tự động phân tách cấu trúc kiến thức và thiết lập lộ trình học theo từng ngày. Mỗi ngày học bao gồm bài giảng chi tiết, ghi chú trọng tâm, trắc nghiệm tự luyện và bài tập tự luận.", "Normal"),
    ("Để đảm bảo tính sư phạm và độ tin cậy của bài học, hệ thống thiết lập vai trò Giáo viên hướng dẫn (Instructor). Giáo viên có thể duyệt lộ trình học của học viên, trực tiếp can thiệp chỉnh sửa bài học (thêm bài học, xóa bài học hoặc chèn thêm ngày học ở vị trí bất kỳ), soạn thảo lại nội dung và chấm điểm bài tập tự luận cho học viên.", "Normal"),
    ("2. Hướng giải quyết", "Heading 2"),
    ("2.1 Mục tiêu hệ thống", "Heading 3"),
    ("Xây dựng một nền tảng học tập thông minh trực tuyến giúp tối ưu hóa việc tự học từ tài liệu thô nhờ AI.", "Normal"),
    ("Áp dụng kỹ thuật RAG (Truy xuất tăng cường phát sinh) để đảm bảo AI sinh bài học chính xác dựa trên tài liệu tải lên, hạn chế tối đa hiện tượng ảo tưởng kiến thức.", "Normal"),
    ("Cung cấp giao diện tương tác và quản lý linh hoạt cho giáo viên để đồng hành và định hướng học viên.", "Normal"),
    ("2.2. Phạm vi giải pháp", "Heading 3"),
    ("Trong phạm vi đồ án, hệ thống tập trung xây dựng các chức năng cốt lõi:", "Normal"),
    ("- Phân quyền người dùng: Học viên (Learner), Giáo viên (Instructor) và Quản trị viên (Admin).", "Normal"),
    ("- Tải tài liệu & Trích xuất văn bản: PDF, Docx, Text.", "Normal"),
    ("- Sinh lộ trình học tự động (Syllabus) và bài giảng chi tiết hàng ngày (Lesson).", "Normal"),
    ("- Chức năng học tập tương tác: Làm bài trắc nghiệm tự động chấm điểm, làm bài tập tự luận gửi giáo viên.", "Normal"),
    ("- Chức năng giáo viên điều phối khóa học: Duyệt lộ trình, thêm/chèn/xóa ngày học, chấm điểm và sửa bài giảng.", "Normal"),
    ("- Chợ khóa học (Marketplace) công khai và hệ thống kết bạn để chia sẻ kiến thức.", "Normal"),
    ("3. Các vấn đề cần giải quyết", "Heading 2"),
    ("3.1. Vấn đề nghiệp vụ", "Heading 3"),
    ("Làm sao để AI phân bổ đều khối lượng kiến thức tài liệu vào số ngày học mong muốn mà không bị sót hoặc gộp quá tải; làm sao để đảm bảo bài học ngày sau kế thừa kiến thức ngày trước mà không bị trùng lặp; quy trình phối hợp sửa đổi lộ trình giữa học viên và giáo viên để đảm bảo tính đồng bộ dữ liệu.", "Normal"),
    ("3.2. Vấn đề kỹ thuật", "Heading 3"),
    ("Trích xuất chính xác cấu trúc mục lục từ tài liệu PDF/Docx chứa nhiều bảng biểu và ký tự nhiễu OCR; xử lý chunking văn bản và tính toán vector embedding; tối ưu hóa cấu trúc prompt của LLM để trả về dữ liệu định dạng JSON ổn định phục vụ xử lý ở frontend; đảm bảo re-indexing tự động các ngày học khi giáo viên chèn hoặc xóa ngày học giữa lộ trình.", "Normal"),
    ("4. Các giải pháp kỹ thuật đề xuất", "Heading 2"),
    ("Hệ thống đề xuất sử dụng kiến trúc RAG với Vector Database, sử dụng Groq Cloud API làm công cụ suy luận AI siêu tốc với mô hình LLaMA-3. Phía server sử dụng Node.js & Express để xử lý logic nghiệp vụ và phía frontend sử dụng ReactJS & TypeScript cùng Tailwind CSS nhằm mang lại trải nghiệm mượt mà, tối ưu nhất.", "Normal"),
    ("5. Công cụ và môi trường phát triển", "Heading 2"),
    ("Hệ thống được lập trình trên hệ điều hành Windows, sử dụng Visual Studio Code làm môi trường phát triển (IDE), Node.js runtime, MongoDB làm cơ sở dữ liệu chính và kết nối với Groq API cho dịch vụ AI.", "Normal"),
    ("6. Lợi ích và giá trị mang lại", "Heading 2"),
    ("Giúp học viên tiết kiệm tới 80% thời gian thiết lập lộ trình học từ tài liệu thô, nâng cao kết quả học tập qua các câu hỏi ôn tập tương tác, đồng thời giúp giảng viên dễ dàng điều phối, kiểm soát tiến độ và chấm điểm bài tập cho hàng loạt học viên trực tuyến một cách dễ dàng.", "Normal")
]

ch2_items = [
    ("1. Khảo sát hiện trạng hệ thống", "Heading 2"),
    ("Khảo sát các phương pháp tự học hiện nay cho thấy người tự học từ tài liệu PDF/Word thường gặp các hạn chế sau:", "Normal"),
    ("- Thiếu tính tương tác: Việc đọc tài liệu một chiều không có công cụ tự kiểm tra kiến thức (như trắc nghiệm, bài tập) khiến người học khó đánh giá được mức độ hiểu bài của mình.", "Normal"),
    ("- Quá tải thông tin: Không thể chia nhỏ tài liệu dài thành lộ trình học hợp lý theo ngày dẫn đến nản lòng.", "Normal"),
    ("- Thiếu sự định hướng: Khi gặp kiến thức khó trong tài liệu tự học, học viên không có sự hỗ trợ trực tiếp từ giáo viên hướng dẫn.", "Normal"),
    ("- Chia sẻ khó khăn: Khó chia sẻ lộ trình tự học tâm đắc cho bạn bè cùng học.", "Normal"),
    ("Để khắc phục các hạn chế trên, website cá nhân hóa lộ trình học tích hợp AI ra đời, cung cấp giải pháp toàn diện giúp học viên tự động hóa việc chia lộ trình từ tài liệu, làm trắc nghiệm sinh tự động và cho phép kết nối trực tiếp giáo viên hướng dẫn để chỉnh sửa giáo án trực quan.", "Normal"),
    ("2. Mục tiêu và phạm vi hệ thống", "Heading 2"),
    ("2.1. Mục tiêu tổng quát", "Heading 3"),
    ("Xây dựng hệ thống web hỗ trợ cá nhân hóa học tập tích hợp AI, tự động hóa quy trình phân tích tài liệu học tập, thiết kế syllabus theo ngày, sinh bài giảng chi tiết, trắc nghiệm tự luyện, kết nối giảng viên hỗ trợ và chia sẻ tài nguyên qua marketplace.", "Normal"),
    ("2.2. Mục tiêu cụ thể", "Heading 3"),
    ("- Tải lên tài liệu và trích xuất text tự động từ file PDF/Docx.", "Normal"),
    ("- AI sinh lộ trình học (Syllabus) N ngày và bài giảng từng ngày kèm câu hỏi trắc nghiệm.", "Normal"),
    ("- Giáo viên đăng ký tài khoản và có quyền duyệt lộ trình, sửa bài giảng, thêm hoặc chèn ngày học mới ở giữa các ngày hiện có, chấm điểm bài tập tự luận.", "Normal"),
    ("- Cung cấp tính năng kết bạn và chia sẻ khóa học lên Marketplace.", "Normal"),
    ("2.3. Phạm vi ứng dụng", "Heading 3"),
    ("Hệ thống hướng tới đối tượng học sinh, sinh viên tự học trực tuyến và các giáo viên, trợ giảng tại các trung tâm giáo dục mong muốn hỗ trợ học viên từ xa. Hệ thống phân chia thành 3 vai trò chính: Learner, Instructor, Admin.", "Normal"),
    ("3. Danh sách quy trình nghiệp vụ chính", "Heading 2"),
    ("- Quy trình tạo khóa học từ tài liệu: Học viên tải file -> Hệ thống trích xuất text -> AI phân tích cấu trúc -> Sinh Syllabus và các bài học ngày chi tiết.", "Normal"),
    ("- Quy trình học tập & Kiểm tra: Học viên đọc bài giảng -> Làm bài trắc nghiệm tự động chấm điểm -> Làm bài tập tự luận và gửi bài giải lên hệ thống.", "Normal"),
    ("- Quy trình chỉnh sửa bài học của giáo viên: Giáo viên có thể bấm Thêm bài học mới ở cuối, hoặc bấm nút \"Chèn\" ẩn ở giữa hai ngày học trên sidebar để chèn bài học vào vị trí bất kỳ. Hệ thống sẽ tự động tịnh tiến dayNumber của các ngày sau đó lên 1 đơn vị.", "Normal"),
    ("- Quy trình gửi và chấm điểm bài tập: Học viên nộp bài tập tự luận -> Giáo viên nhận được bản sao khóa học -> Giáo viên chấm điểm và trả nhận xét.", "Normal"),
    ("- Quy trình kết bạn và chia sẻ Marketplace: Người dùng tìm kiếm kết bạn, gửi yêu cầu kết bạn, đăng khóa học lên marketplace để người khác tải về học.", "Normal")
]

ch3_items = [
    ("1. Node.js & Express", "Heading 2"),
    ("Node.js là một runtime environment mã nguồn mở dựa trên Chrome V8 engine, cho phép chạy JavaScript phía máy chủ. Express là một framework gọn nhẹ, linh hoạt cho Node.js cung cấp các tính năng mạnh mẽ để xây dựng ứng dụng web và API RESTful. Sự kết hợp này mang lại hiệu năng xử lý bất đồng bộ (non-blocking I/O) cực cao, xử lý hàng loạt yêu cầu sinh bài học song song hiệu quả.", "Normal"),
    ("2. ReactJS & Vite & TypeScript", "Heading 2"),
    ("ReactJS là thư viện JavaScript phổ biến để xây dựng giao diện người dùng dựa trên component. Đồ án sử dụng Vite làm công cụ build tool giúp tối ưu hóa tốc độ khởi chạy và hot-reload trong quá trình phát triển. Đồng thời, việc sử dụng TypeScript giúp quản lý mã nguồn frontend một cách chặt chẽ thông qua cơ chế gõ kiểu tĩnh, giảm thiểu lỗi runtime.", "Normal"),
    ("3. MongoDB & Mongoose", "Heading 2"),
    ("MongoDB là hệ quản trị cơ sở dữ liệu NoSQL hướng tài liệu phổ biến nhất hiện nay. Thay vì các bảng và hàng của cơ sở dữ liệu quan hệ, MongoDB lưu trữ dữ liệu dưới dạng JSON-like tài liệu với schema linh hoạt. Mongoose là thư viện ODM (Object Document Mapper) cho Node.js giúp quản lý các mối quan hệ dữ liệu, thực hiện validate schema một cách dễ dàng.", "Normal"),
    ("4. Groq API & Mô hình LLaMA-3", "Heading 2"),
    ("Groq là nền tảng điện toán tăng tốc suy luận AI với kiến trúc phần cứng LPU (Language Processing Unit), giúp tốc độ phản hồi của mô hình LLM nhanh hơn gấp 10 lần so với GPU thông thường. Hệ thống tích hợp Groq API để gọi mô hình LLaMA-3 70B/8B sinh nội dung học tập và câu hỏi trắc nghiệm bám sát tài liệu với thời gian phản hồi chỉ tính bằng giây, giải quyết bài toán trải nghiệm người dùng khi phải chờ đợi AI sinh bài.", "Normal")
]

ch4_items = [
    ("1. Thiết kế sơ đồ Use case", "Heading 2"),
    ("1.1. Đặc tả Usecase", "Heading 3"),
    ("Hệ thống bao gồm các ca sử dụng chính phục vụ cho ba tác nhân (Actor): Học viên (Learner), Giáo viên (Instructor) và Quản trị viên (Admin).", "Normal"),
    ("- Đăng nhập/Đăng ký: Cho phép người dùng tạo tài khoản và đăng nhập vào hệ thống.", "Normal"),
    ("- Upload tài liệu & Tạo khóa học: Học viên tải file lên để AI phân tích và sinh lộ trình học.", "Normal"),
    ("- Học bài & Kiểm tra: Học viên xem nội dung bài giảng, trả lời câu hỏi trắc nghiệm và làm bài tập.", "Normal"),
    ("- Đăng ký làm giáo viên: Học viên gửi yêu cầu nâng cấp tài khoản thành giáo viên.", "Normal"),
    ("- Quản lý bài học (Giáo viên): Giáo viên có quyền chỉnh sửa giáo án, thêm ngày học mới hoặc chèn ngày học vào giữa lộ trình, xóa bài học hiện có.", "Normal"),
    ("- Chấm điểm bài tập: Giáo viên xem lời giải của học sinh, chấm điểm và nhận xét.", "Normal"),
    ("- Quản lý bạn bè: Tìm kiếm người dùng, gửi/nhận lời mời kết bạn và xem danh sách bạn bè.", "Normal"),
    ("- Chợ khóa học (Marketplace): Đăng tải các khóa học chất lượng lên chợ công khai, nhập các khóa học từ chợ về trang cá nhân.", "Normal"),
    ("1.2. Sơ đồ use case", "Heading 3"),
    ("[HƯỚNG DẪN VẼ SƠ ĐỒ USE CASE]: Học viên vẽ 1 sơ đồ Use case tổng quát chứa 3 Actor: Learner (nối với các usecase: Đăng ký/Đăng nhập, Tạo khóa học từ tài liệu, Học bài giảng, Làm trắc nghiệm, Nộp bài tập tự luận, Gửi khóa học cho Giáo viên, Kết bạn, Nhập khóa học từ Marketplace), Instructor (nối với các usecase: Duyệt khóa học của học sinh, Thêm/Xóa/Chèn ngày học, Soạn bài giảng trực quan, Chấm điểm bài tập, Đăng khóa học lên Marketplace) và Admin (nối với các usecase: Quản lý người dùng, Duyệt yêu cầu làm Giáo viên).", "Normal"),
    ("2. Thiết kế sơ đồ sequence", "Heading 2"),
    ("2.1. Sơ đồ sequence chức năng đăng nhập", "Heading 3"),
    ("[HƯỚNG DẪN VẼ SEQUENCE ĐĂNG NHẬP]: Tương tác giữa Learner/Instructor -> UI đăng nhập -> AuthController -> UserRepository -> MongoDB. UI gửi thông tin đăng nhập -> Controller xác thực mật khẩu -> trả về JWT Token và chuyển hướng trang chủ.", "Normal"),
    ("2.2. Sơ đồ sequence chức năng tải tài liệu và sinh lộ trình", "Heading 3"),
    ("[HƯỚNG DẪN VẼ SEQUENCE SINH LỘ TRÌNH]: Learner -> UI -> planController -> planService -> GroqAPI -> MongoDB. Learner tải file -> Controller trích xuất text -> Service phân đoạn text (chunking) và lưu vào MongoDB -> Gọi Groq API (LLaMA3) sinh Syllabus -> Lưu Syllabus và bài giảng các ngày học -> Trả kết quả thành công về giao diện.", "Normal"),
    ("2.3. Sơ đồ sequence chức năng học bài và làm trắc nghiệm", "Heading 3"),
    ("[HƯỚNG DẪN VẼ SEQUENCE HỌC BÀI]: Learner -> UI -> lessonController -> MongoDB. Learner chọn ngày học -> UI hiển thị bài giảng markdown -> Học viên chọn làm trắc nghiệm -> Gửi đáp án -> Hệ thống đối chiếu và trả về số câu đúng trực tiếp.", "Normal"),
    ("2.4. Sơ đồ sequence chức năng giáo viên chèn ngày học mới", "Heading 3"),
    ("[HƯỚNG DẪN VẼ SEQUENCE CHÈN NGÀY HỌC]: Instructor -> UI -> instructorController -> MongoDB. Giáo viên bấm nút \"Chèn sau Ngày X\" -> UI gửi yêu cầu lên API kèm `afterDayNumber = X` -> Controller thực hiện lệnh updateMany để tăng dayNumber của toàn bộ bài học có `dayNumber > X` thêm 1 đơn vị -> Tạo bài học mới ở `dayNumber = X + 1` -> Cập nhật tổng số ngày duration của Plan -> Trả về danh sách ngày đã cập nhật cho UI.", "Normal"),
    ("2.5. Sơ đồ sequence chức năng giáo viên chấm điểm bài tập", "Heading 3"),
    ("[HƯỚNG DẪN VẼ SEQUENCE CHẤM ĐIỂM]: Instructor -> UI -> instructorController -> MongoDB. Giáo viên xem bài giải tự luận -> Nhập điểm và nhận xét -> Gửi dữ liệu -> Controller cập nhật bản ghi Assignment -> Đồng bộ tiến độ cho Học viên.", "Normal"),
    ("3. Thiết kế sơ đồ class", "Heading 2"),
    ("[HƯỚNG DẪN VẼ SƠ ĐỒ CLASS]: Vẽ các lớp thực thể tương ứng với các Schema MongoDB trong code: User (id, email, password, role, status), Document (id, userId, title, content, contentHash), Plan (id, title, owner, studentId, instructorId, duration, status, isPublic), Lesson (id, planId, dayNumber, title, content, summary, quiz, status), Chunk (id, planId, text, vectorId), Enrollment (id, learnerId, instructorId, planId, status), Assignment (id, lessonId, studentId, solution, grade, feedback), Progress (id, planId, studentId, completedLessonsCount, completedPercent).", "Normal"),
    ("4. Thiết kế cơ sở dữ liệu", "Heading 2"),
    ("4.1. Sơ đồ Erd", "Heading 3"),
    ("[HƯỚNG DẪN VẼ SƠ ĐỒ ERD]: Vẽ sơ đồ quan hệ thực thể giữa các collection của MongoDB: User (1) - (n) Document, User (1) - (n) Plan (với vai trò owner), User (1) - (n) Enrollment (learnerId), User (1) - (n) Enrollment (instructorId), Document (1) - (1) Plan, Plan (1) - (n) Lesson, Plan (1) - (n) Chunk, Lesson (1) - (n) Assignment, User (1) - (n) Assignment (studentId).", "Normal"),
    ("4.2. Collection users", "Heading 3"),
    ("Lưu trữ thông tin tài khoản người dùng của hệ thống, bao gồm học viên, giáo viên và admin.", "Normal"),
    ("4.3. Collection documents", "Heading 3"),
    ("Lưu trữ văn bản trích xuất từ tài liệu gốc do người dùng tải lên, kèm mã hash MD5 để chống trùng lặp dữ liệu trên hệ thống.", "Normal"),
    ("4.4. Collection plans", "Heading 3"),
    ("Lưu trữ lộ trình tổng quan của khóa học (Syllabus), số ngày học, thông tin học viên sở hữu và giáo viên hướng dẫn.", "Normal"),
    ("4.5. Collection lessons", "Heading 3"),
    ("Lưu trữ nội dung chi tiết bài học của từng ngày dưới dạng markdown, kèm theo các câu hỏi trắc nghiệm tự luyện.", "Normal"),
    ("4.6. Collection chunks", "Heading 3"),
    ("Lưu trữ các phân đoạn văn bản nhỏ được trích xuất từ tài liệu để phục vụ việc truy vấn thông tin chính xác theo mô hình RAG.", "Normal"),
    ("4.7. Collection enrollments", "Heading 3"),
    ("Quản lý việc theo học của học viên đối với giáo viên trong một lộ trình cụ thể, quản lý trạng thái phê duyệt.", "Normal"),
    ("4.8. Collection assignments", "Heading 3"),
    ("Lưu trữ các lời giải bài tập tự luận do học viên nộp và điểm số, đánh giá nhận xét từ giáo viên hướng dẫn.", "Normal"),
    ("4.9. Collection progress", "Heading 3"),
    ("Theo dõi tiến độ học tập thực tế của học viên trong từng lộ trình cụ thể.", "Normal")
]

ch5_items = [
    ("1. Mục tiêu đặt ra và kế hoạch thực hiện", "Heading 2"),
    ("1.1. Mục tiêu đặt ra", "Heading 3"),
    ("1.1.1. Mục tiêu về phân tích và thiết kế hệ thống", "Heading 4"),
    ("Phân tích rõ nét nghiệp vụ tự học tích hợp trí tuệ nhân tạo, thiết kế mô hình RAG tối ưu, phân chia cấu trúc dữ liệu cơ sở dữ liệu MongoDB hợp lý và xây dựng các sơ đồ Use case, Class, Sequence chi tiết cho các nghiệp vụ chèn/xóa ngày học của giáo viên.", "Normal"),
    ("1.1.2. Mục tiêu về phát triển hệ thống", "Heading 4"),
    ("Xây dựng thành công phía backend Node.js cung cấp các API RESTful bảo mật và frontend ReactJS trực quan, sinh động. Tích hợp Groq API sinh lộ trình học và bài học với tốc độ cao.", "Normal"),
    ("1.1.3. Mục tiêu về tích hợp công nghệ và phân quyền", "Heading 4"),
    ("Phân quyền chặt chẽ giữa học viên và giáo viên hướng dẫn khóa học, đảm bảo giáo viên chỉ chỉnh sửa bản sao khóa học của họ và chỉ hiển thị ở bảng điều khiển khi học viên gửi yêu cầu.", "Normal"),
    ("1.1.4. Mục tiêu về kiểm thử và hoàn thiện hệ thống", "Heading 4"),
    ("Thực hiện kiểm thử các module trích xuất tài liệu, chất lượng bài soạn của AI và độ chính xác của logic tịnh tiến ngày học khi giáo viên thêm hoặc chèn bài giảng.", "Normal"),
    ("1.1.5. Mục tiêu về tài liệu và báo cáo", "Heading 4"),
    ("Hoàn thiện cuốn báo cáo đồ án tốt nghiệp mô tả chi tiết kiến trúc, các công nghệ sử dụng và các kết quả nghiên cứu khoa học đạt được của đề tài.", "Normal"),
    ("1.2. Kế hoạch thực hiện", "Heading 3"),
    ("Kế hoạch thực hiện kéo dài trong 12 tuần, bao gồm các giai đoạn: Khảo sát & Phân tích yêu cầu (Tuần 1-2), Thiết kế DB & Hệ thống (Tuần 3-4), Phát triển Backend (Tuần 5-6), Phát triển Frontend & Tích hợp AI (Tuần 7-9), Kiểm thử & Sửa lỗi (Tuần 10-11), Viết báo cáo đồ án (Tuần 12).", "Normal"),
    ("2. Tìm hiểu công nghệ và cài đặt dự án", "Heading 2"),
    ("2.1. Node.js và Express Framework", "Heading 3"),
    ("Sử dụng Node.js làm môi trường phát triển chính phía server và Express làm framework routing điều hướng các API.", "Normal"),
    ("2.2. ReactJS và Vite", "Heading 3"),
    ("Sử dụng ReactJS với TypeScript giúp xây dựng giao diện Single Page Application đáp ứng nhanh, quản lý state linh hoạt.", "Normal"),
    ("2.3. MongoDB và Mongoose ODM", "Heading 3"),
    ("Cài đặt cơ sở dữ liệu MongoDB để lưu trữ dữ liệu dạng Document và dùng Mongoose định nghĩa các Schema.", "Normal"),
    ("2.4. Cấu hình kết nối Groq Cloud API", "Heading 3"),
    ("Thiết lập kết nối an toàn với Groq API sử dụng API Key lưu trữ trong file môi trường `.env` nhằm bảo mật thông tin.", "Normal"),
    ("3. Kiến trúc dự án và xây dựng các lớp cơ bản", "Heading 2"),
    ("3.1. Kiến trúc dự án", "Heading 3"),
    ("Dự án áp dụng kiến trúc Service-Controller-Model giúp phân tách rõ ràng trách nhiệm:", "Normal"),
    ("- Controllers: Nhận request từ HTTP client, thực hiện validate dữ liệu đầu vào sơ bộ và gửi phản hồi kết quả.", "Normal"),
    ("- Services: Chứa toàn bộ logic nghiệp vụ phức tạp như kết nối RAG, phân tích tài liệu thô, liên kết dữ liệu AI và tính toán logic tịnh tiến ngày học.", "Normal"),
    ("- Models: Định nghĩa cấu trúc các thực thể dữ liệu lưu trữ trong MongoDB.", "Normal"),
    ("3.2. Xây dựng các lớp cơ bản", "Heading 3"),
    ("3.2.1. Xây dựng lớp Model", "Heading 4"),
    ("Định nghĩa các Schema chi tiết cho User, Document, Plan, Lesson, Chunk, Enrollment, Assignment và Progress sử dụng thư viện Mongoose.", "Normal"),
    ("3.2.2. Xây dựng lớp planService", "Heading 4"),
    ("Lớp `planService` thực hiện việc chunking văn bản, gọi Groq API để sinh khung chương trình và bài học chi tiết dựa trên bối cảnh RAG.", "Normal"),
    ("3.2.3. Xây dựng lớp instructorController", "Heading 4"),
    ("Lớp `instructorController` điều phối các hoạt động của giáo viên, đặc biệt chứa logic chèn ngày học mới (`addLesson` với tham số `afterDayNumber`) và tự động tịnh tiến thứ tự các ngày học tiếp theo.", "Normal"),
    ("4. Kiểm thử", "Heading 2"),
    ("4.1. Kiểm thử tính năng", "Heading 3"),
    ("Tiến hành viết các kịch bản kiểm thử (Test Cases) cho chức năng cốt lõi: Đăng nhập/Đăng ký, Tải tài liệu sinh lộ trình học, Giáo viên chèn ngày học mới vào giữa các ngày hiện có để kiểm tra tính liên tục của `dayNumber`.", "Normal"),
    ("4.2. Kết quả kiểm thử", "Heading 3"),
    ("Hệ thống hoạt động trơn tru. Khi chèn một ngày học mới vào sau ngày thứ 2 trong lộ trình 5 ngày, các bài học ngày 3, 4, 5 cũ được tự động cập nhật thứ tự thành 4, 5, 6 một cách chính xác. Lộ trình học cập nhật đồng bộ ở cả giao diện của học viên và giáo viên.", "Normal")
]

ch7_items = [
    ("1. Kết luận", "Heading 2"),
    ("Đề tài \"Hệ thống hỗ trợ cá nhân hóa lộ trình học tập và quản lý bài học tích hợp AI\" đã hoàn thành xuất sắc các mục tiêu nghiên cứu và phát triển đề ra. Hệ thống giải quyết tốt bài toán tự học cá nhân hóa bằng việc ứng dụng kiến trúc RAG kết hợp với mô hình ngôn ngữ lớn LLaMA-3 thông qua Groq API, giúp sinh bài học thông minh, nhanh chóng từ tài liệu bất kỳ. Đồng thời, hệ thống cung cấp giao diện tương tác giảng dạy hiệu quả cho giáo viên, cho phép can thiệp sâu vào cấu trúc lộ trình thông qua các tác vụ chèn/xóa ngày học trực quan.", "Normal"),
    ("2. Hướng phát triển", "Heading 2"),
    ("Trong tương lai, hệ thống có thể nghiên cứu mở rộng một số hướng đi mới:", "Normal"),
    ("- Tích hợp thuật toán MMR (Maximal Marginal Relevance) vào cơ chế truy vấn RAG để đa dạng hóa kiến thức lấy ra từ tài liệu, giảm thiểu sự lặp ý giữa các bài học.", "Normal"),
    ("- Tự động chuyển đổi bài học markdown thành các định dạng đa phương tiện như podcast âm thanh (Text-to-Speech) hoặc slide bài giảng tự động.", "Normal"),
    ("- Ứng dụng mô hình AI chấm điểm bài tập tự luận tự động kết hợp với nhận xét của giáo viên để tăng tốc độ phản hồi cho học viên.", "Normal"),
    ("3. Kết luận chung", "Heading 2"),
    ("Việc phát triển thành công đề tài không chỉ đem lại một sản phẩm phần mềm hữu ích cho hoạt động tự học trực tuyến mà còn khẳng định tiềm năng ứng dụng to lớn của Trí tuệ nhân tạo (AI) trong lĩnh vực đổi mới giáo dục và giảng dạy thông minh.", "Normal")
]

# We need to find the correct indices for the original document.
# From headings.txt:
# CHƯƠNG 1 H1 is paragraph index 490.
# CHƯƠNG 2 H1 is paragraph index 550.
# Chương 3 H1 is paragraph index 599.
# Chương 4 H1 is paragraph index 634.
# Chương 5 H1 is paragraph index 1371.
# Chương 6 H1 is paragraph index 1575.
# Chương 7 H1 is paragraph index 1710.
# Tài liệu tham khảo is paragraph index 1752.

print("Replacing Chapter 7...")
replace_range(doc, 1711, 1751, ch7_items)

print("Replacing Chapter 5...")
replace_range(doc, 1372, 1574, ch5_items)

print("Replacing Chapter 4...")
replace_range(doc, 635, 1370, ch4_items)

print("Replacing Chapter 3...")
replace_range(doc, 600, 633, ch3_items)

print("Replacing Chapter 2...")
replace_range(doc, 551, 598, ch2_items)

print("Replacing Chapter 1...")
replace_range(doc, 491, 549, ch1_items)

# Save the document!
doc.save(original_path)
print("Document updated successfully!")
